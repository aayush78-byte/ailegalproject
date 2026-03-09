import io
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract

def extract_text(file, file_extension):
    """
    Extract text from uploaded file (in-memory processing)
    
    Args:
        file: FileStorage object from Flask
        file_extension: str (.pdf, .docx, .doc)
    
    Returns:
        str: Extracted text
    """
    try:
        file_bytes = file.read()
        
        if file_extension == '.pdf':
            return extract_from_pdf(file_bytes)
        elif file_extension in ['.docx', '.doc']:
            return extract_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        raise

def extract_from_pdf(file_bytes):
    """
    Extract text from PDF using PyMuPDF
    Falls back to OCR if text extraction fails
    """
    text = ""
    
    try:
        # Open PDF from bytes
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            
            # If page has little/no text, try OCR
            if len(page_text.strip()) < 50:
                page_text = extract_with_ocr(page)
            
            text += page_text + "\n\n"
        
        pdf_document.close()
        
        return text.strip()
    
    except Exception as e:
        print(f"PDF extraction failed: {str(e)}")
        # Try OCR as last resort
        return extract_pdf_with_ocr(file_bytes)

def extract_from_docx(file_bytes):
    """
    Extract text from DOCX file
    """
    try:
        # Open DOCX from bytes
        doc = Document(io.BytesIO(file_bytes))
        
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        print(f"DOCX extraction failed: {str(e)}")
        raise

def extract_with_ocr(page):
    """
    Extract text from a PDF page using OCR
    
    Args:
        page: fitz.Page object
    
    Returns:
        str: OCR extracted text
    """
    try:
        # Render page to image
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
        img_bytes = pix.tobytes("png")
        
        # Convert to PIL Image
        image = Image.open(io.BytesIO(img_bytes))
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang='eng')
        
        return text
    
    except Exception as e:
        print(f"OCR failed: {str(e)}")
        return ""

def extract_pdf_with_ocr(file_bytes):
    """
    Extract text from entire PDF using OCR (fallback method)
    """
    try:
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text += extract_with_ocr(page) + "\n\n"
        
        pdf_document.close()
        return text.strip()
    
    except Exception as e:
        print(f"PDF OCR extraction failed: {str(e)}")
        return ""

def clean_extracted_text(text):
    """
    Clean and normalize extracted text
    """
    # Remove excessive whitespace
    text = ' '.join(text.split())
    
    # Remove common artifacts
    text = text.replace('\x00', '')
    text = text.replace('\uf0b7', '•')
    
    return text